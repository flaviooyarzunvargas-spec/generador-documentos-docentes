from fastapi import FastAPI, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.openapi.utils import get_openapi
from pydantic import BaseModel
from docx import Document
from docx.shared import Inches
from pathlib import Path
import unicodedata
import uuid
import re

from generadores.funciones import (
    generar_grafico_funcion,
    generar_grafico_sistema
)

from generadores.estadistica import (
    generar_grafico_barras,
    generar_grafico_sectores,
    generar_grafico_histograma,
    generar_grafico_dispersion
)

from generadores.geometria import (
    generar_grafico_circunferencia,
    generar_grafico_triangulo,
    generar_grafico_poligono,
    generar_grafico_traslacion
)


app = FastAPI(
    title="Generador de Documentos Docentes",
    version="0.1.0"
)

OUTPUT_DIR = Path("/tmp/documentos")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

app.mount(
    "/documentos",
    StaticFiles(directory=str(OUTPUT_DIR)),
    name="documentos"
)


class DocumentoRequest(BaseModel):
    tipo: str
    titulo: str
    curso: str
    asignatura: str
    contenido: str


class DocumentoResponse(BaseModel):
    mensaje: str
    titulo: str
    curso: str
    asignatura: str
    tipo: str
    archivo: str
    url: str


def custom_openapi():
    if app.openapi_schema:
        return app.openapi_schema

    openapi_schema = get_openapi(
        title="Generador de Documentos Docentes",
        version="0.1.0",
        description="API para generar documentos docentes en formato DOCX.",
        routes=app.routes,
    )

    openapi_schema["servers"] = [
        {
            "url": "https://generador-documentos-docentes.onrender.com",
            "description": "Servidor en Render"
        }
    ]

    app.openapi_schema = openapi_schema
    return app.openapi_schema


app.openapi = custom_openapi


def normalizar_texto(texto: str) -> str:
    texto = texto.lower().strip()
    texto = unicodedata.normalize("NFD", texto)
    texto = "".join(
        c for c in texto
        if unicodedata.category(c) != "Mn"
    )
    return texto


@app.get("/")
def inicio():
    return {"mensaje": "API generadora de documentos activa"}


@app.get("/test")
def test():
    return {"estado": "ok"}


@app.get("/ping")
def ping():
    return {"ping": "pong"}


def seleccionar_plantilla(tipo: str):
    base_dir = Path(__file__).parent
    plantillas_dir = base_dir / "plantillas"

    tipo_normalizado = normalizar_texto(tipo)

    plantillas = {
        "guia": "GuíaICA.docx",
        "prueba": "PruebaICA.docx",
        "planificacion": "PlanificacionICA.docx",
        "planificacion clase": "PlanificacionICA.docx",
        "rubrica": "RubricaICA.docx",
        "solucionario": "GuíaICA.docx"
    }

    nombre_plantilla = plantillas.get(tipo_normalizado)

    if not nombre_plantilla:
        raise HTTPException(
            status_code=400,
            detail=(
                f"Tipo de documento no válido: {tipo}. "
                "Use guia, prueba, planificacion, rubrica o solucionario."
            )
        )

    ruta_plantilla = plantillas_dir / nombre_plantilla

    if not ruta_plantilla.exists():
        raise HTTPException(
            status_code=404,
            detail=f"No existe la plantilla: {nombre_plantilla}"
        )

    return ruta_plantilla


def parsear_parametros(contenido: str) -> dict:
    parametros = {}
    partes = contenido.split(";")

    for parte in partes:
        if "=" in parte:
            clave, valor = parte.split("=", 1)
            parametros[clave.strip()] = valor.strip()

    return parametros


def procesar_linea_con_grafico_funcion(doc: Document, linea: str) -> bool:
    coincidencia = re.search(r"\[GRAFICO_FUNCION:(.*?)\]", linea)
    if not coincidencia:
        return False

    parametros = parsear_parametros(coincidencia.group(1))
    expresion = parametros.get("expr")
    titulo = parametros.get("titulo", f"Gráfico de {expresion}")
    x_min = float(parametros.get("x_min", -5))
    x_max = float(parametros.get("x_max", 5))

    if not expresion:
        doc.add_paragraph("Error: marcador de gráfico sin expresión.")
        return True

    ruta_grafico = generar_grafico_funcion(
        expresion, titulo, x_min, x_max, OUTPUT_DIR
    )
    doc.add_picture(str(ruta_grafico), width=Inches(5.5))
    return True


def procesar_linea_con_grafico_sistema(doc: Document, linea: str) -> bool:
    coincidencia = re.search(r"\[GRAFICO_SISTEMA:(.*?)\]", linea)
    if not coincidencia:
        return False

    parametros = parsear_parametros(coincidencia.group(1))
    expresion_f = parametros.get("f")
    expresion_g = parametros.get("g")
    titulo = parametros.get("titulo", "Sistema de ecuaciones")
    x_min = float(parametros.get("x_min", -5))
    x_max = float(parametros.get("x_max", 5))

    if not expresion_f or not expresion_g:
        doc.add_paragraph("Error: marcador de sistema sin funciones f o g.")
        return True

    ruta_grafico = generar_grafico_sistema(
        expresion_f, expresion_g, titulo, x_min, x_max, OUTPUT_DIR
    )
    doc.add_picture(str(ruta_grafico), width=Inches(5.5))
    return True


def procesar_linea_con_grafico_barras(doc: Document, linea: str) -> bool:
    coincidencia = re.search(r"\[GRAFICO_BARRAS:(.*?)\]", linea)
    if not coincidencia:
        return False

    parametros = parsear_parametros(coincidencia.group(1))
    titulo = parametros.get("titulo", "Gráfico de barras")
    categorias_texto = parametros.get("categorias")
    valores_texto = parametros.get("valores")
    etiqueta_x = parametros.get("etiqueta_x", "Categoría")
    etiqueta_y = parametros.get("etiqueta_y", "Valor")

    if not categorias_texto or not valores_texto:
        doc.add_paragraph("Error: marcador de barras sin categorías o valores.")
        return True

    categorias = [c.strip() for c in categorias_texto.split(",")]

    try:
        valores = [float(v.strip()) for v in valores_texto.split(",")]
    except ValueError:
        doc.add_paragraph("Error: los valores del gráfico de barras deben ser numéricos.")
        return True

    ruta_grafico = generar_grafico_barras(
        categorias, valores, titulo, etiqueta_x, etiqueta_y, OUTPUT_DIR
    )
    doc.add_picture(str(ruta_grafico), width=Inches(5.5))
    return True


def procesar_linea_con_grafico_sectores(doc: Document, linea: str) -> bool:
    coincidencia = re.search(r"\[GRAFICO_SECTORES:(.*?)\]", linea)
    if not coincidencia:
        return False

    parametros = parsear_parametros(coincidencia.group(1))
    titulo = parametros.get("titulo", "Gráfico circular")
    categorias_texto = parametros.get("categorias")
    valores_texto = parametros.get("valores")

    if not categorias_texto or not valores_texto:
        doc.add_paragraph("Error: marcador de sectores sin categorías o valores.")
        return True

    categorias = [c.strip() for c in categorias_texto.split(",")]

    try:
        valores = [float(v.strip()) for v in valores_texto.split(",")]
    except ValueError:
        doc.add_paragraph("Error: los valores del gráfico de sectores deben ser numéricos.")
        return True

    try:
        ruta_grafico = generar_grafico_sectores(
            categorias, valores, titulo, OUTPUT_DIR
        )
    except Exception as e:
        doc.add_paragraph(f"Error al generar gráfico de sectores: {e}")
        return True

    doc.add_picture(str(ruta_grafico), width=Inches(5.5))
    return True


def procesar_linea_con_grafico_histograma(doc: Document, linea: str) -> bool:
    coincidencia = re.search(r"\[GRAFICO_HISTOGRAMA:(.*?)\]", linea)
    if not coincidencia:
        return False

    parametros = parsear_parametros(coincidencia.group(1))
    titulo = parametros.get("titulo", "Histograma")
    datos_texto = parametros.get("datos")
    etiqueta_x = parametros.get("etiqueta_x", "Datos")
    etiqueta_y = parametros.get("etiqueta_y", "Frecuencia")

    try:
        bins = int(parametros.get("bins", 5))
    except ValueError:
        bins = 5

    if not datos_texto:
        doc.add_paragraph("Error: marcador de histograma sin datos.")
        return True

    try:
        datos = [float(valor.strip()) for valor in datos_texto.split(",")]
    except ValueError:
        doc.add_paragraph("Error: los datos del histograma deben ser numéricos.")
        return True

    try:
        ruta_grafico = generar_grafico_histograma(
            datos, titulo, etiqueta_x, etiqueta_y, bins, OUTPUT_DIR
        )
    except Exception as e:
        doc.add_paragraph(f"Error al generar histograma: {e}")
        return True

    doc.add_picture(str(ruta_grafico), width=Inches(5.5))
    return True


def procesar_linea_con_grafico_dispersion(doc: Document, linea: str) -> bool:
    coincidencia = re.search(r"\[GRAFICO_DISPERSION:(.*?)\]", linea)
    if not coincidencia:
        return False

    parametros = parsear_parametros(coincidencia.group(1))
    titulo = parametros.get("titulo", "Diagrama de dispersión")
    x_texto = parametros.get("x")
    y_texto = parametros.get("y")
    etiqueta_x = parametros.get("etiqueta_x", "X")
    etiqueta_y = parametros.get("etiqueta_y", "Y")

    if not x_texto or not y_texto:
        doc.add_paragraph("Error: marcador de dispersión sin valores x o y.")
        return True

    try:
        valores_x = [float(v.strip()) for v in x_texto.split(",")]
        valores_y = [float(v.strip()) for v in y_texto.split(",")]
    except ValueError:
        doc.add_paragraph("Error: los valores de dispersión deben ser numéricos.")
        return True

    try:
        ruta_grafico = generar_grafico_dispersion(
            valores_x, valores_y, titulo, etiqueta_x, etiqueta_y, OUTPUT_DIR
        )
    except Exception as e:
        doc.add_paragraph(f"Error al generar dispersión: {e}")
        return True

    doc.add_picture(str(ruta_grafico), width=Inches(5.5))
    return True


def procesar_linea_con_grafico_circunferencia(doc: Document, linea: str) -> bool:
    coincidencia = re.search(r"\[GRAFICO_CIRCUNFERENCIA:(.*?)\]", linea)
    if not coincidencia:
        return False

    parametros = parsear_parametros(coincidencia.group(1))
    centro_texto = parametros.get("centro", "0,0")
    radio_texto = parametros.get("radio")
    titulo = parametros.get("titulo", "Circunferencia")

    if not radio_texto:
        doc.add_paragraph("Error: marcador de circunferencia sin radio.")
        return True

    try:
        centro_partes = [p.strip() for p in centro_texto.split(",")]
        centro_x = float(centro_partes[0])
        centro_y = float(centro_partes[1])
        radio = float(radio_texto)

        x_min = float(parametros.get("x_min", centro_x - radio - 2))
        x_max = float(parametros.get("x_max", centro_x + radio + 2))
        y_min = float(parametros.get("y_min", centro_y - radio - 2))
        y_max = float(parametros.get("y_max", centro_y + radio + 2))

    except Exception:
        doc.add_paragraph(
            "Error: formato inválido en marcador de circunferencia. "
            "Use centro=0,0; radio=3."
        )
        return True

    ruta_grafico = generar_grafico_circunferencia(
        centro_x, centro_y, radio, titulo, x_min, x_max, y_min, y_max, OUTPUT_DIR
    )
    doc.add_picture(str(ruta_grafico), width=Inches(5.5))
    return True


def procesar_linea_con_grafico_triangulo(doc: Document, linea: str) -> bool:
    coincidencia = re.search(r"\[GRAFICO_TRIANGULO:(.*?)\]", linea)
    if not coincidencia:
        return False

    parametros = parsear_parametros(coincidencia.group(1))
    puntos_texto = parametros.get("puntos")
    etiquetas_texto = parametros.get("etiquetas", "A,B,C")
    titulo = parametros.get("titulo", "Triángulo")

    if not puntos_texto:
        doc.add_paragraph("Error: marcador de triángulo sin puntos.")
        return True

    try:
        puntos = []
        for par in puntos_texto.split("|"):
            x_texto, y_texto = par.split(",")
            puntos.append((float(x_texto.strip()), float(y_texto.strip())))

        etiquetas = [e.strip() for e in etiquetas_texto.split(",")]
        xs = [p[0] for p in puntos]
        ys = [p[1] for p in puntos]

        x_min = float(parametros.get("x_min", min(xs) - 1))
        x_max = float(parametros.get("x_max", max(xs) + 1))
        y_min = float(parametros.get("y_min", min(ys) - 1))
        y_max = float(parametros.get("y_max", max(ys) + 1))

    except Exception:
        doc.add_paragraph(
            "Error: formato inválido en marcador de triángulo. "
            "Use puntos=0,0|4,0|0,3; etiquetas=A,B,C."
        )
        return True

    try:
        ruta_grafico = generar_grafico_triangulo(
            puntos, etiquetas, titulo, x_min, x_max, y_min, y_max, OUTPUT_DIR
        )
    except Exception as e:
        doc.add_paragraph(f"Error al generar triángulo: {e}")
        return True

    doc.add_picture(str(ruta_grafico), width=Inches(5.5))
    return True


def procesar_linea_con_grafico_poligono(doc: Document, linea: str) -> bool:
    coincidencia = re.search(r"\[GRAFICO_POLIGONO:(.*?)\]", linea)
    if not coincidencia:
        return False

    parametros = parsear_parametros(coincidencia.group(1))
    puntos_texto = parametros.get("puntos")
    etiquetas_texto = parametros.get("etiquetas", "")
    titulo = parametros.get("titulo", "Polígono")

    if not puntos_texto:
        doc.add_paragraph("Error: marcador de polígono sin puntos.")
        return True

    try:
        puntos = []
        for par in puntos_texto.split("|"):
            x_texto, y_texto = par.split(",")
            puntos.append((float(x_texto.strip()), float(y_texto.strip())))

        etiquetas = [
            e.strip()
            for e in etiquetas_texto.split(",")
            if e.strip()
        ]

        xs = [p[0] for p in puntos]
        ys = [p[1] for p in puntos]

        x_min = float(parametros.get("x_min", min(xs) - 1))
        x_max = float(parametros.get("x_max", max(xs) + 1))
        y_min = float(parametros.get("y_min", min(ys) - 1))
        y_max = float(parametros.get("y_max", max(ys) + 1))

    except Exception:
        doc.add_paragraph(
            "Error: formato inválido en marcador de polígono. "
            "Use puntos=0,0|4,0|5,3; etiquetas=A,B,C."
        )
        return True

    try:
        ruta_grafico = generar_grafico_poligono(
            puntos, etiquetas, titulo, x_min, x_max, y_min, y_max, OUTPUT_DIR
        )
    except Exception as e:
        doc.add_paragraph(f"Error al generar polígono: {e}")
        return True

    doc.add_picture(str(ruta_grafico), width=Inches(5.5))
    return True


def procesar_linea_con_grafico_traslacion(doc: Document, linea: str) -> bool:
    coincidencia = re.search(r"\[GRAFICO_TRASLACION:(.*?)\]", linea)
    if not coincidencia:
        return False

    parametros = parsear_parametros(coincidencia.group(1))
    puntos_texto = parametros.get("puntos")
    etiquetas_texto = parametros.get("etiquetas", "")
    titulo = parametros.get("titulo", "Traslación")

    if not puntos_texto:
        doc.add_paragraph("Error: marcador de traslación sin puntos.")
        return True

    try:
        dx = float(parametros.get("dx", 0))
        dy = float(parametros.get("dy", 0))

        puntos = []
        for par in puntos_texto.split("|"):
            x_texto, y_texto = par.split(",")
            puntos.append((float(x_texto.strip()), float(y_texto.strip())))

        etiquetas = [
            e.strip()
            for e in etiquetas_texto.split(",")
            if e.strip()
        ]

        puntos_trasladados = [
            (x + dx, y + dy)
            for x, y in puntos
        ]

        todos_x = [p[0] for p in puntos] + [p[0] for p in puntos_trasladados]
        todos_y = [p[1] for p in puntos] + [p[1] for p in puntos_trasladados]

        x_min = float(parametros.get("x_min", min(todos_x) - 1))
        x_max = float(parametros.get("x_max", max(todos_x) + 1))
        y_min = float(parametros.get("y_min", min(todos_y) - 1))
        y_max = float(parametros.get("y_max", max(todos_y) + 1))

    except Exception:
        doc.add_paragraph(
            "Error: formato inválido en marcador de traslación. "
            "Use puntos=0,0|4,0|2,3; dx=3; dy=2; etiquetas=A,B,C."
        )
        return True

    try:
        ruta_grafico = generar_grafico_traslacion(
            puntos,
            etiquetas,
            dx,
            dy,
            titulo,
            x_min,
            x_max,
            y_min,
            y_max,
            OUTPUT_DIR
        )
    except Exception as e:
        doc.add_paragraph(f"Error al generar traslación: {e}")
        return True

    doc.add_picture(str(ruta_grafico), width=Inches(5.5))
    return True


def procesar_linea_con_elementos_visuales(doc: Document, linea: str) -> bool:
    if procesar_linea_con_grafico_funcion(doc, linea):
        return True

    if procesar_linea_con_grafico_sistema(doc, linea):
        return True

    if procesar_linea_con_grafico_barras(doc, linea):
        return True

    if procesar_linea_con_grafico_sectores(doc, linea):
        return True

    if procesar_linea_con_grafico_histograma(doc, linea):
        return True

    if procesar_linea_con_grafico_dispersion(doc, linea):
        return True

    if procesar_linea_con_grafico_circunferencia(doc, linea):
        return True

    if procesar_linea_con_grafico_triangulo(doc, linea):
        return True

    if procesar_linea_con_grafico_poligono(doc, linea):
        return True

    if procesar_linea_con_grafico_traslacion(doc, linea):
        return True

    return False


@app.post(
    "/generar-documento",
    response_model=DocumentoResponse
)
def generar_documento(data: DocumentoRequest):
    try:
        plantilla = seleccionar_plantilla(data.tipo)
        doc = Document(str(plantilla))

        doc.add_heading(data.titulo, level=1)
        doc.add_paragraph(f"Curso: {data.curso}")
        doc.add_paragraph(f"Asignatura: {data.asignatura}")
        doc.add_paragraph(f"Tipo de documento: {data.tipo}")

        doc.add_heading("Contenido", level=2)

        for linea in data.contenido.split("\n"):
            linea_limpia = linea.strip()

            if procesar_linea_con_elementos_visuales(doc, linea_limpia):
                continue

            doc.add_paragraph(linea)

        nombre_archivo = f"{uuid.uuid4()}.docx"
        ruta_salida = OUTPUT_DIR / nombre_archivo

        doc.save(str(ruta_salida))

        url_archivo = (
            "https://generador-documentos-docentes.onrender.com/"
            f"documentos/{nombre_archivo}"
        )

        return DocumentoResponse(
            mensaje="Documento generado correctamente",
            titulo=data.titulo,
            curso=data.curso,
            asignatura=data.asignatura,
            tipo=data.tipo,
            archivo=nombre_archivo,
            url=url_archivo
        )

    except HTTPException:
        raise

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=str(e)
        )
